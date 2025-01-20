from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import numpy as np
from datetime import datetime

class DocumentProcessor:
    def __init__(self):
        self.feedback_templates = {
            "商业计划": self._analyze_business_plan,
            "市场分析": self._analyze_market_research,
            "财务预测": self._analyze_financial_forecast,
            "团队介绍": self._analyze_team_profile
        }
        
    def analyze_uploaded_documents(self, documents):
        """分析上传的文档并生成反馈"""
        analysis_results = {}
        for doc_type, content in documents.items():
            if doc_type in self.feedback_templates:
                analysis_results[doc_type] = self.feedback_templates[doc_type](content)
                
        return self._generate_comprehensive_feedback(analysis_results)
    
    def generate_solution_document(self, consultation_data):
        """生成解决方案文档"""
        doc = Document()
        self._add_document_header(doc)
        self._add_executive_summary(doc, consultation_data)
        self._add_problem_analysis(doc, consultation_data)
        self._add_solution_details(doc, consultation_data)
        self._add_implementation_plan(doc, consultation_data)
        self._add_appendix(doc, consultation_data)
        
        return doc
        
    def _analyze_business_plan(self, content):
        """分析商业计划书"""
        return {
            "完整性评分": self._evaluate_completeness(content),
            "可行性评分": self._evaluate_feasibility(content),
            "创新性评分": self._evaluate_innovation(content),
            "改进建议": self._generate_improvement_suggestions(content)
        }
        
    def _analyze_market_research(self, content):
        """分析市场研究报告"""
        return {
            "市场洞察": self._extract_market_insights(content),
            "竞争分析": self._analyze_competition(content),
            "机会点": self._identify_opportunities(content),
            "风险点": self._identify_risks(content)
        }
        
    def _generate_comprehensive_feedback(self, analysis_results):
        """生成综合反馈意见"""
        return {
            "总体评价": self._calculate_overall_score(analysis_results),
            "关键发现": self._summarize_key_findings(analysis_results),
            "改进建议": self._compile_improvement_suggestions(analysis_results),
            "下一步行动": self._suggest_next_steps(analysis_results)
        }
        
    def _add_document_header(self, doc):
        """添加文档头部"""
        title = doc.add_heading('创业项目解决方案', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加基本信息表格
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        cells = [
            ("项目名称", ""),
            ("咨询日期", datetime.now().strftime("%Y-%m-%d")),
            ("版本号", "V1.0"),
            ("文档状态", "终稿")
        ]
        for i, (key, value) in enumerate(cells):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value
            
    def _add_executive_summary(self, doc, data):
        """添加执行摘要"""
        doc.add_heading('执行摘要', level=1)
        summary = doc.add_paragraph()
        summary.add_run('项目背景：').bold = True
        summary.add_run(data.get('background', ''))
        
        summary = doc.add_paragraph()
        summary.add_run('核心问题：').bold = True
        summary.add_run(data.get('core_issues', ''))
        
        summary = doc.add_paragraph()
        summary.add_run('解决方案概述：').bold = True
        summary.add_run(data.get('solution_summary', ''))
        
    def _add_implementation_plan(self, doc, data):
        """添加实施计划"""
        doc.add_heading('实施计划', level=1)
        
        # 添加时间线
        timeline = data.get('timeline', [])
        table = doc.add_table(rows=len(timeline)+1, cols=3)
        table.style = 'Table Grid'
        
        # 添加表头
        headers = table.rows[0].cells
        headers[0].text = '时间节点'
        headers[1].text = '关键任务'
        headers[2].text = '预期成果'
        
        # 添加内容
        for i, item in enumerate(timeline, 1):
            row = table.rows[i].cells
            row[0].text = item.get('time', '')
            row[1].text = item.get('task', '')
            row[2].text = item.get('expected_outcome', '') 